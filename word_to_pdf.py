from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

def convert_word_to_pdf(word_data):
    """
    Convert Word document content to PDF
    Args:
        word_data: BytesIO object containing the Word document
    Returns:
        str: Path to the generated PDF file
    """
    # Create a temporary directory if it doesn't exist
    if not os.path.exists('temp'):
        os.makedirs('temp')
    
    # Save the Word document temporarily
    temp_word_path = 'temp/temp_doc.docx'
    with open(temp_word_path, 'wb') as f:
        f.write(word_data.getvalue())
    
    # Read the Word document
    doc = Document(temp_word_path)
    
    # Create PDF
    pdf_path = 'temp/generated_word_report.pdf'
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    
    # Set font and starting position
    c.setFont("Helvetica", 12)
    y = height - 50  # Start from top with margin
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        if y < 50:  # If we're near the bottom of the page
            c.showPage()  # Start a new page
            c.setFont("Helvetica", 12)
            y = height - 50
        
        # Split text into lines that fit the page width
        words = paragraph.text.split()
        line = []
        for word in words:
            line.append(word)
            if c.stringWidth(' '.join(line), "Helvetica", 12) > width - 100:
                line.pop()
                c.drawString(50, y, ' '.join(line))
                y -= 20
                line = [word]
        
        if line:
            c.drawString(50, y, ' '.join(line))
            y -= 20
        
        y -= 10  # Add space between paragraphs
    
    # Process tables
    for table in doc.tables:
        if y < 100:  # If we're near the bottom of the page
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 50
        
        # Draw table
        cell_width = (width - 100) / len(table.columns)
        cell_height = 20
        
        # Draw headers
        x = 50
        for cell in table.rows[0].cells:
            c.rect(x, y, cell_width, cell_height)
            c.drawString(x + 5, y + 5, cell.text)
            x += cell_width
        
        y -= cell_height
        
        # Draw data rows
        for row in table.rows[1:]:
            if y < 50:
                c.showPage()
                c.setFont("Helvetica", 12)
                y = height - 50
                x = 50
            
            x = 50
            for cell in row.cells:
                c.rect(x, y, cell_width, cell_height)
                c.drawString(x + 5, y + 5, cell.text)
                x += cell_width
            y -= cell_height
        
        y -= 20  # Add space after table
    
    c.save()
    
    # Clean up temporary Word file
    os.remove(temp_word_path)
    
    return pdf_path 